#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Sun Jan 19 12:35:02 2020

@author: Steve Satterfield
"""

import constant
from PyPDF4 import PdfFileReader
from docx import Document
from textblob import TextBlob
from course import Course
from knowledgearea import KnowledgeArea
import string
from paragraph import Paragraph
import nltk
from nltk.stem.porter import PorterStemmer
from firebase import firebase


firstPositionOfCKA = 0
currentKnowledgeArea = KnowledgeArea()

pdfFileName = 'Fall2019_LLFullCatalog.pdf'


def getTextPDF(pdfFileName, password = ''):
    pdf_file = open(pdfFileName, 'rb')
    read_pdf = PdfFileReader(pdf_file)
    
    if password != '':
        read_pdf.decrypt(password)
        
    text = []
    for i in range(0, read_pdf.getNumPages()-1):
        text.append(read_pdf.getPage(i).extractText())
        
    #return '\n'.join(text)
    return text

    
    
def getMatchingKnowledgeArea(candidate):
    # find matching knowledge area and return Id
    
    for knowledgearea in knowledgeAreas:
        if knowledgearea.getText().strip() == str(candidate).strip():
            return knowledgearea
        
    return KnowledgeArea()


def ParagraphIsKnowledgeArea(pdocument, p, i, lastPos, pText):
    # make the determination that the content of a paragraph consists of a Knowledge Area
    
    global firstPositionOfCKA
    
    if p.style.name == constant.STYLE_KNOWLEDGEAREA and pText.strip() and pText[lastPos-1].isdigit():
        pPrev = document.paragraphs[i - 1]
        pPrevText = pPrev.text.strip()
        pPrevLasPos = len(pPrevText)
        
        pPrevRunCount = len(pPrev.runs)
        
        pPrevIsBold = 'false'
        if pPrevRunCount > 0:
            pPrevIsBold = 'true' if pPrev.runs[0].bold else 'false' 
            
        if pPrevIsBold == 'false':
            if pPrevRunCount > 1:
                pPrevIsBold = 'true' if pPrev.runs[1].bold else 'false' 
        
        
        if pPrevText != "":
            if pPrev.style.name == constant.STYLE_NORMAL and pPrevIsBold == 'true' and not (pPrevText[pPrevLasPos-1].isdigit()):
                if firstPositionOfCKA == 0:
                    firstPositionOfCKA = i
        
                return 'true', pPrevText
            
    return 'false', ""

       
def ExtractKnowledgeAreas(firebase, document, knowledgeAreas):
    # find the knowledge areas in the document and write them to a database table
    
    global firstPositionOfCKA
    firstPositionOfCKA = 0
    
    for i, p in enumerate(document.paragraphs):
        pText = p.text.strip()
        lastPos = len(pText)
        kaText = ""
        
        isKnowledgeArea, kaText = ParagraphIsKnowledgeArea(document, p, i, lastPos, pText)
        
        
        if isKnowledgeArea == 'true' and kaText:
            knowledgeArea = KnowledgeArea()
            knowledgeArea.setText(kaText)
            knowledgeAreas.append(knowledgeArea)
                    
                    
        diff = (i - firstPositionOfCKA)
        if firstPositionOfCKA > 0 and diff > 120:
            break;            
                
    for knowledgeArea in knowledgeAreas:
        newKnowledgeArea = {
            'Content':knowledgeArea.getText()    
        }
        result = firebase.post('KnowledgeArea',newKnowledgeArea)
        knowledgeArea.setId(result.get("name"))
        
      
def ExtractCourseAndDescription(firebase, document, knowledgeAreas, courses):
    # build a list of courses with foreign key back to knowledge area and insert into DB
    
    global currentKnowledgeArea
    knowledgeAreaId = ""
    knowledgeAreaTitle = ""
    candidateId = ""
    partialTitle = []
    fulltitle = ""
   
    for p in document.paragraphs:
        if p.style.name == constant.STYLE_NORMAL:
            currentKnowledgeArea = getMatchingKnowledgeArea(p.text.strip())
            candidateId = currentKnowledgeArea.getId()
            
            if candidateId != "":
                knowledgeAreaId = str(candidateId)
                knowledgeAreaTitle = currentKnowledgeArea.getText()
                #print("knowledgeAreaId:{}".format(knowledgeAreaId))
                  
        else:
            knowledgeAreaId = ""
    
       
        if knowledgeAreaId != "" and p.text.strip() != knowledgeAreaTitle and p.text.strip() != "":
            pText = p.text.strip()
            lastPos = len(pText)

            if not pText[lastPos-1].isdigit():
                partialTitle.append(p.text.strip())
            else:
                if len(partialTitle) > 0:
                    partialTitle.append(p.text.strip())
                    fulltitle = ' '.join(partialTitle)
                else:
                    fulltitle = pText
                    
                  
                fulltitle = fulltitle.replace("*","")
                    
                course = Course()
                course.setKnowledgeAreaId(knowledgeAreaId)
                course.setKnowledgeArea(currentKnowledgeArea.getText())
                course.setTOCEntry(p.text.strip())
                course.setTitle(fulltitle.rstrip(string.digits).strip())
                courses.append(course)
                partialTitle = []

    ExtractCourseDescriptions(document, courses)   
    
    
    for course in courses:
        if course.getDescription() == "":
            pNumber = 0
            takeDescription = 0
            descriptionElements = []
            
            for p in document.paragraphs:
                pNumber += 1
                if p.style.name == constant.STYLE_NORMAL and p.text != "" and p.runs[0].font.name == "Calibri" and p.runs[0].font.bold:
                    pText = p.text.replace("New! ","")
                    if p.text.__contains__(course.getTitle()) or course.getTitle().__contains__(pText):
                        print("takeDescription is on")
                        takeDescription = 1
                    else:
                        if takeDescription == 1 and p.text.strip().__contains__("Sessions"):
                            takeDescription = 1
                        if takeDescription == 1 and p.text.strip().__contains__("Magic"):
                            takeDescription = 1
                        else:
                            takeDescription = 0
                        
                if p.style.name == constant.STYLE_BODY_TEXT and takeDescription == 1:
                    descriptionElements.append(p.text.strip())
                    
            fullDescription = course.getDescription()
            
            if len(descriptionElements) > 0:
                fullDescription = ' '.join(descriptionElements)
                
            course.setDescription(fullDescription)
            
                
    
    for course in courses:
        if course.getKnowledgeArea() == course.getTitle():
            continue
        
        newCourse = {
            'KnowledgeAreaId':course.getKnowledgeAreaId(),
            'Name':course.getTitle(),
            'Description':course.getDescription()
        }
        result = firebase.post('Course',newCourse)
        course.setId(result)


def ExtractCourses(firebase, document, knowledgeAreas, courses):
    # build a list of courses with foreign key back to knowledge area and insert into DB
    
    global currentKnowledgeArea
    knowledgeAreaId = ""
    knowledgeAreaTitle = ""
    candidateId = ""
    partialTitle = []
    fulltitle = ""
   
    for p in document.paragraphs:
        if p.style.name == constant.STYLE_NORMAL:
            currentKnowledgeArea = getMatchingKnowledgeArea(p.text.strip())
            candidateId = currentKnowledgeArea.getId()
            
            if candidateId != "":
                knowledgeAreaId = str(candidateId)
                knowledgeAreaTitle = currentKnowledgeArea.getText()
                #print("knowledgeAreaId:{}".format(knowledgeAreaId))
                  
        else:
            knowledgeAreaId = ""
    
       
        if knowledgeAreaId != "" and p.text.strip() != knowledgeAreaTitle and p.text.strip() != "":
            pText = p.text.strip()
            lastPos = len(pText)

            if not pText[lastPos-1].isdigit():
                partialTitle.append(p.text.strip())
            else:
                if len(partialTitle) > 0:
                    partialTitle.append(p.text.strip())
                    fulltitle = ' '.join(partialTitle)
                else:
                    fulltitle = pText
                    
                  
                fulltitle = fulltitle.replace("*","")
                    
                course = Course()
                course.setKnowledgeAreaId(knowledgeAreaId)
                course.setKnowledgeArea(currentKnowledgeArea.getText())
                course.setTOCEntry(p.text.strip())
                course.setTitle(fulltitle.rstrip(string.digits).strip())
                courses.append(course)
                partialTitle = []

    ExtractCourseDescriptions(document, courses)   
    
    for course in courses:
        if course.getKnowledgeArea() == course.getTitle():
            continue
        
        newCourse = {
            'KnowledgeAreaId':course.getKnowledgeAreaId(),
            'Name':course.getTitle(),
            'Description':course.getDescription()
        }
        result = firebase.post('Course',newCourse)
        course.setId(result)
        


def HeaderIsCourseTitle(candidate):
    # compare candidate with every course title until match found or eol    
    
    for course in courses:
        #print("course title: {}".format(course.getTitle()))
        
        compareThis = []
        courseTitle = course.getTitle().strip()
        lastPos = len(courseTitle)
        #lastChar = courseTitle.charAt(lastPos-1)
        
        
        if courseTitle != "":
            
            if courseTitle.__contains__(":") and courseTitle[lastPos-1] != ":":
                (key,value) = courseTitle.split(':',1)
                compareThis.append(key.strip())
                compareThis.append(value.strip())
                compareThis.append(courseTitle)
            else:
                compareThis.append(courseTitle)
            
       
            if any(ele in candidate for ele in compareThis) and len(candidate) >= len(courseTitle):
                return 'true'
        
    return 'false'    
        
def GetAssociatedCourse(courses,candidate):
    # compare candidate with every course title until match found or eol    
    
    
    candidate = candidate.replace("New! ","")
    
        
    for course in courses:
        #print("course title: {}".format(course.getTitle()))
        
        compareThis = []
        courseTitle = course.getTitle().strip()
        lastPos = len(courseTitle)
        #lastChar = courseTitle.charAt(lastPos-1)
        
        
        if courseTitle != "":
            
            if courseTitle.__contains__(":") and courseTitle[lastPos-1] != ":":
                (key,value) = courseTitle.split(':',1)
                compareThis.append(key.strip())
                compareThis.append(value.strip())
                compareThis.append(courseTitle)
            else:
                compareThis.append(courseTitle)
            
       
            #if any(ele in candidate for ele in compareThis) and len(candidate) >= len(courseTitle):
            if any(ele in candidate for ele in compareThis):
                return course
        
    nullCourse = Course()
    nullCourse.setKnowledgeArea = "No Match"
    nullCourse.setTOCEntry = "No Match"
    nullCourse.setTitle = "No Match"
    return nullCourse    
    

def ExtractCourseDescriptions(document, courses):
    
    courseDescriptionOn = 'false'
    allparagraphs = []
    pertinentParagraphs = []
    number = 0
    pHeader = ""
    headerIsCourseTitle = 'false'
    partialTitle = []
    fulltitle = ""
    
    for p in document.paragraphs:
        
        
        if p.text.strip() != "":
            
            pText = p.text.strip()
            lastPos = len(pText)

            if pText[lastPos-1] == ":":
                partialTitle.append(pText)
            else:
                if len(partialTitle) > 0:
                    partialTitle.append(p.text.strip())
                    fulltitle = ' '.join(partialTitle)
                else:
                    fulltitle = pText
                
            number = number + 1
            paragraph = Paragraph()
            paragraph.setNumber(number)
            paragraph.setStyle(p.style.name)
            paragraph.setText(p.text.strip().rstrip(string.digits))
            allparagraphs.append(paragraph)
            
            
    for i, item in enumerate(allparagraphs):
        if item.getStyle() == constant.STYLE_BODY_TEXT:
            pPrev = allparagraphs[i-1]
            pHeader = ""
            if pPrev.getStyle() == constant.STYLE_NORMAL:
                number = number + 1
                p2 = Paragraph()
                p2.setNumber(number)
                p2.setStyle(pPrev.getStyle())
                p2.setText(pPrev.getText())
                paragraphs2.append(p2)
                pHeader = pPrev.getText().strip()
                if pHeader != "":
                    headerIsCourseTitle = HeaderIsCourseTitle(pHeader)
                    
            if headerIsCourseTitle == 'true':
                number = number + 1
                pc = Paragraph()
                pc.setNumber(number)
                pc.setStyle(item.getStyle())
                pc.setHeader(pHeader)
                pc.setText(item.getText())
                pertinentParagraphs.append(pc)
            
    
    associatedCourse = Course()  
    currentCourseDescription = "" 
    courseTitlePosition = -1
     
    
    for i, p in enumerate(pertinentParagraphs):
        courseTitleCandidate = p.getHeader().strip()
        #print("courseTitleCandidate: {}".format(courseTitleCandidate))
        if p.getHeader().strip() != "" and courseDescriptionOn == 'true' and i > courseTitlePosition:
            associatedCourse.setDescription(currentCourseDescription)
            courseDescriptionOn = 'false'
            currentCourseDescription = ""
            courseTitlePosition = -1
        if p.getHeader().strip() != "" and courseDescriptionOn == 'false':
            courseDescriptionOn = 'true'
            associatedCourse = GetAssociatedCourse(courses,p.getHeader().strip())
            #print("return from getAssociatedCourse(): {}".format(associatedCourse.getTitle()))
            courseTitlePosition = i
            currentCourseDescription += p.getText().strip()
        if courseDescriptionOn == 'true' and i > courseTitlePosition:
            currentCourseDescription += p.getText().strip()


def SentenceIsLO(sentence):
    # determine if the sentence meets the criteria for a learning outcome
    
    nullSentence = ""
    
    for word, pos in sentence.tags:
        if constant.LL_VERBS.count(word) >= 1 and pos in [constant.VERB, constant.VERB_GERUND, constant.VERB_PAST_PARTICIPLE, constant.VERB_PAST_TENSE, constant.VERB_SINGULAR_PRESENT, constant.VERB_THIRD_PERSON_SINGULAR]:
            return 'true', sentence
        if constant.LL_NOUNS.count(word) >= 1 and pos in [constant.NOUN_PLURAL, constant.NOUN_PROPER_PLURAL, constant.NOUN_PROPER_SINGULAR, constant.NOUN_SINGULAR]:
            return 'true', sentence
        if constant.LL_ADJECTIVES.count(word) >= 1 and pos in [constant.ADJECTIVE, constant.ADJECTIVE_COMPARATIVE, constant.ADJECTIVE_SUPERLATIVE]:
            return 'true', sentence
        for phrase in constant.LL_PHRASES:
            if sentence.noun_phrases.count(phrase) >= 1:
                return 'true', sentence
            
        return 'false', nullSentence
    

def ExtractLearningOutcomes(firebase, document, courses):
    # load each course description into blob and get the sentences
    # tag parts of speech in each sentence
    # compare verb in each sentence to list of Bloom action verbs
    # if match, select sentence as a learning outcome
    
    
    for course in courses:
        
       
        blob = TextBlob(course.getDescription())
        for sentence in blob.sentences:
            isLO, candidateLO = SentenceIsLO(sentence)
            
            if isLO == 'true':
                newLO = {
                    'CourseId':course.getId(),
                    'Text':candidateLO
                }
                result = firebase.post('LearningObjective', newLO)
                
   
def printPDF(pdfFileName, password):
    print(getTextPDF(pdfFileName, password))
    
   
def loadPDFIntoDocx(pdfFileName, password):
    strDoc = getTextPDF(pdfFileName, password)
    
    doc = docx.Document('demo.docx')
    
def loadPDFPagesIntoTextBlob(pdfFileName, password):
    text = []
    needyCourses = []
    
    for course in courses:
        if course.getDescription() == "":
            needyCourses.append(course.getTitle())
            
    text = getTextPDF(pdfFileName, password)
    
    for item in text:
        # for title in needyCourses:
        #     if str(item).__contains__(title):
        #         print(item)
                
        # blob = TextBlob(item)
        # for sentence in blob.sentences:
            
        #     for course in courses:
        #         if str(course.getTitle()).__contains__(str(sentence.strip())):
        #              print("course: {} contains: {}".format(course.getTitle(), str(sentence)))
        
        if str(item).__contains__("Charcoal Drawing"):
            print(str(item))
        
    
      
def loadPDFIntoTextBlob(pdfFileName, password):
    
    text = []
    text = getTextPDF(pdfFileName, password)

    #print("{}".format(text))
    
    blob = TextBlob(str(text))
    
    for sentence in blob.sentences:
        print("{}".format(str(sentence)))
    
    # for course in courses:
        
    #     print(course.getTitle())
    #     for sentence in blob.sentences:
    #         if sentence.strip() == course.getTitle():
    #             print("match: {}".format(str(sentence)))
    #         elif str(course.getTitle()).__contains__(str(sentence.strip())):
    #             print("contains:{}".format(str(sentence)))
    #         else:
    #             if sentence.find(str(course.getTitle())):
    #                 print("found: {} in sentence: {}".format(course.getTitle(), str(sentence)))
                
                 
firebase = firebase.FirebaseApplication('https://scholacity-org-test.firebaseio.com/')
#document = Document('Fall2019_LLFullCatalog.docx')

        
knowledgeAreas = []
courses = []
paragraphs1 = []
paragraphs2 = []



#ExtractKnowledgeAreas(firebase, document, knowledgeAreas)
#ExtractCourseAndDescription(firebase, document, knowledgeAreas, courses)
#ExtractLearningOutcomes(firebase, document, courses)

loadPDFIntoTextBlob(pdfFileName, "")


        
    
