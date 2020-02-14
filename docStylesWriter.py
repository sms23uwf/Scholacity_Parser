#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Mon Jan 20 15:23:01 2020

@author: claudiasatterfield
"""

import csv
import constant
from docx import Document
from textblob import TextBlob
from course import Course
from paragraph import Paragraph
import string
import math
        
allCourseFile = "Courses.csv"
allCourseFile2 = "CoursesWithDescriptions.csv"


paragraphs1 = []
paragraphs2 = []

number = 0
courses = []
coursesWithDescriptions = []

currentKnowledgeArea = ""
document = Document('Fall2019_LLFullCatalog.docx')

def writeCourses():
      
    fieldNamesCourse = ['course title']    
    global courses
    global currentKnowledgeArea
    
    for p in document.paragraphs:
        #print(p.text.strip())
        if p.style.name == constant.STYLE_NORMAL:
            if p.text.strip() == constant.KNOWLEDGE_AREA_1:
                currentKnowledgeArea = constant.KNOWLEDGE_AREA_1
            elif p.text.strip() == constant.KNOWLEDGE_AREA_2:
                currentKnowledgeArea = constant.KNOWLEDGE_AREA_2
            elif p.text.strip() == constant.KNOWLEDGE_AREA_3:
                currentKnowledgeArea = constant.KNOWLEDGE_AREA_3
            elif p.text.strip() == constant.KNOWLEDGE_AREA_4:
                currentKnowledgeArea = constant.KNOWLEDGE_AREA_4
            elif p.text.strip() == constant.KNOWLEDGE_AREA_5:
                currentKnowledgeArea = constant.KNOWLEDGE_AREA_5
            elif p.text.strip() == constant.KNOWLEDGE_AREA_6:
                currentKnowledgeArea = constant.KNOWLEDGE_AREA_6
            elif p.text.strip() == constant.KNOWLEDGE_AREA_7:
                currentKnowledgeArea = constant.KNOWLEDGE_AREA_7
            else:
                currentKnowledgeArea = currentKnowledgeArea
        else:
            currentKnowledgeArea = ""
    
       
        if currentKnowledgeArea != "":
            course = Course()
            course.setKnowledgeArea(currentKnowledgeArea)
            course.setTitle(p.text.strip().rstrip(string.digits))
            courses.append(course)

    with open(allCourseFile, "w") as allCourseOutput:
        allCourseWriter = csv.writer(allCourseOutput)
    
        for course in courses:
            if course.getKnowledgeArea() != course.getTitle():
                allCourseWriter.writerow([course.getKnowledgeArea(), course.getTitle(), course.getDescription()])

    currentKnowledgeArea = ""
    currentCourse = ""
    
    for course in courses:
        currentKnowledgeArea = ""
        currentCourse = ""
        courseDescription = ""
        
        for p in document.paragraphs:
            if p.style.name == constant.STYLE_NORMAL and p.text.strip().__contains__(course.getTitle()):
                currentKnowledgeArea = course.getKnowledgeArea()
                currentCourse = course.getTitle()
            elif p.style.name == constant.STYLE_BODY_TEXT and currentCourse != "":
                currentKnowledgeArea = currentKnowledgeArea
                currentCourse = currentCourse
                courseDescription += p.text.strip()
            else:
               if currentCourse != "":
                   course = Course()
                   course.setKnowledgeArea(currentKnowledgeArea)
                   course.setTitle(currentCourse)
                   course.setDescription(courseDescription)
                   coursesWithDescriptions.append(course)
               
                   currentKnowledgeArea = ""
                   currentCourse = ""

                
    with open(allCourseFile2, "w") as allCourseDescriptionOutput:
        allCourseDescriptionWriter = csv.writer(allCourseDescriptionOutput)
    
        for course in coursesWithDescriptions:
            if course.getKnowledgeArea() != course.getTitle():
                allCourseDescriptionWriter.writerow([course.getKnowledgeArea(), course.getTitle(), course.getDescription()])
        
        
        
def writeDocStyles():
    global number
    writeFileRuns = "runs.csv"
    writeFile = "docStyles6.csv"
    #document = Document('Fall2019_LLFullCatalog.docx')
    
    number = 0
    fieldnames = ['number','style','text']
    runs = []

    with open(writeFileRuns, 'w') as outputRuns:
        writerRuns = csv.DictWriter(outputRuns, fieldnames=fieldnames)
        writerRuns.writeheader()
        for r in document.runs:
            number = number + 1
            run = Paragraph()
            run.setNumber(number)
            run.setStyle(r.style.name)
            run.setText(r.text.strip())
            runs.append(run)
            
        for run in runs:
            writerRuns.writerow({run.getNumber(), run.getStyle(), run.getText()})
            
            

    
    with open(writeFile, 'w') as output:
        #fieldnames = ['number','style','text']
        writer = csv.DictWriter(output, fieldnames=fieldnames)
    
        writer.writeheader()
        for p in document.paragraphs:
            if p.text.strip() != "":
                number = number + 1
                paragraph = Paragraph()
                paragraph.setNumber(number)
                paragraph.setStyle(p.style.name)
                paragraph.setText(p.text.replace("New!","").strip().rstrip(string.digits))
                paragraphs1.append(paragraph)
                
        number = 0 
        i = 0
        
        for i, item in enumerate(paragraphs1):
            if item.getStyle() == constant.STYLE_BODY_TEXT:
                pPrev = paragraphs1[i-1]
                if pPrev.getStyle() == constant.STYLE_NORMAL:
                    number = number + 1
                    p2 = Paragraph()
                    p2.setNumber(number)
                    p2.setStyle(pPrev.getStyle())
                    p2.setText(pPrev.getText())
                    paragraphs2.append(p2)
                number = number + 1
                pc = Paragraph()
                pc.setNumber(number)
                pc.setStyle(item.getStyle())
                pc.setText(item.getText())
                paragraphs2.append(pc)
                
                
        for paragraph in paragraphs2:
            writer.writerow({'number': paragraph.getNumber(), 'style': paragraph.getStyle(), 'text': paragraph.getText()})    


writeDocStyles()
