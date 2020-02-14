#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Mon Jan 20 21:18:29 2020

@author: claudiasatterfield
"""

import csv
import constant
from docx import Document
from textblob import TextBlob
from course import Course
from paragraph import Paragraph
from paragraph import Run
import string

allSectionsFile = "Sections.csv"

paragraphs = []
paragraphs1 = []
paragraphs2 = []

number = 0
courses = []

currentKnowledgeArea = ""
document = Document('Fall2019_LLFullCatalog.docx')

def writeSections():
    
    with open(allSectionsFile, "w") as output:
        allSectionsWriter = csv.writer(output)
        for s in document.sections:
            allSectionsWriter.writerow([s.start_type, s.header])
            
        
def writeDocStyles():
    global number
    writeFile = "docStyles01262020.csv"
    writeFileRaw = "docStylesRaw01262020.csv"
    
    #document = Document('Fall2019_LLFullCatalog.docx')
    
    with open(writeFile, 'w') as output:
        fieldnames = ['number','style','text']
        writer = csv.DictWriter(output, fieldnames=fieldnames)
        
    
        writer.writeheader()
        for p in document.paragraphs:
            number = number + 1
            p1 = Paragraph()
            p1.setNumber(number)
            p1.setStyle(p.style.name)
            p1.setStyleType(p.style.type)
            p1.setStyleId(p.style.style_id)
            p1.setFormat(p.paragraph_format)
            p1.setText(p.text.strip())
            p1.setRunCount(len(p.runs))
            for r in p.runs:
                run = Run()
                run.setFont(r.font)
                run.setBold(r.bold)
                run.setItalic(r.italic)
                p1.setRuns(run)
            paragraphs1.append(p1)
                
        number = 0 
        i = 0
        
        for i, item in enumerate(paragraphs1):
            if item.getStyle() == constant.STYLE_BODY_TEXT:
                pPrev = paragraphs1[i-1]
                pPrevminus1 = paragraphs1[i-2]
                if pPrev.getStyle() == constant.STYLE_NORMAL and pPrev:
                    number = number + 1
                    p2 = Paragraph()
                    p2.setNumber(number)
                    p2.setStyle(pPrev.getStyle())
                    p2.setText(pPrev.getText())
                    p2.setFormat(pPrev.getFormat())
                    paragraphs2.append(p2)
                number = number + 1
                pc = Paragraph()
                pc.setNumber(number)
                pc.setStyle(item.getStyle())
                pc.setText(item.getText())
                paragraphs2.append(pc)
                
        for paragraph in paragraphs2:
            writer.writerow({'number': paragraph.getNumber(), 'style': paragraph.getStyle(), 'text': paragraph.getText()})    
    
    with open(writeFileRaw, "w") as outputRaw:
        fieldnames = ['number','style','styleType','styleId', 'format', 'runs', 'bold1', 'bold2', 'text']
        writer = csv.DictWriter(outputRaw, fieldnames=fieldnames)
            
        for paragraph in paragraphs1:
            firstRun = Run()
            firstRun.setBold("nada")
            firstRun.setItalic("nada")
            firstRun.setFont("nada")
            
            secondRun = Run()
            secondRun.setBold("nada")
            secondRun.setItalic("nada")
            secondRun.setFont("nada")
            
            if paragraph.getRunCount() > 0:
                firstRun = paragraph.getRuns()[0]
                
            if paragraph.getRunCount() > 1:
                secondRun = paragraph.getRuns()[1]
                
            writer.writerow({'number': paragraph.getNumber(), 'style': paragraph.getStyle(), 'styleType': paragraph.getStyleType(), 'styleId': paragraph.getStyleId(), 'format': paragraph.getFormat(), 'runs': paragraph.getRunCount(), 'bold1': firstRun.getBold(), 'bold2': secondRun.getBold(), 'text': paragraph.getText()})    

#writeSections()


writeDocStyles()
