"""
    parse_catalog_into_database.py
    --------------------------------

    This module contains a set of functions that implement the extraction
    of the Knowledge Areas, Courses, and stated Learning Outcomes from a
    Course Catalog and inserts them into a Database.


"""

__author__ = "Steven M. Satterfield"

import string
from typing import List, Tuple
import json

from docx import Document
from textblob import TextBlob
from firebase import firebase

from catalog import Catalog
from knowledgearea import KnowledgeArea
from course import Course
from learningobjective import LearningObjective
from learningobjective_course import LearningObjective_course
from paragraph import Paragraph
import constant
import time

existingCatalogRecords = []
existingKnowledgeAreas = []
firstPositionOfCKA = 0
currentKnowledgeArea = KnowledgeArea()
courseNameStyles = [constant.STYLE_NORMAL, constant.STYLE_HEADING]

firebase = firebase.FirebaseApplication('https://scholacity-org.firebaseio.com/')

document = Document('../Catalogs/Fall2019_LLFullCatalog.docx')
documentName = "Fall2019_LLFullCatalog.docx"
documentSemester = "Fall"
documentYear = "2019"

# document = Document('Spring2020_LeisureLearningCatalogFULL.docx')
# documentName = "Spring2020_LeisureLearningCatalogFULL.docx"
# documentSemester = "Spring"
# documentYear = "2020"


knowledgeAreas = []
courses = []
learningObjectives = []
paragraphs1 = []
paragraphs2 = []


def getMatchingKnowledgeArea(candidate: str) -> KnowledgeArea:
    """
        Find matching knowledge area and return the Knowledge Area as a string.

        :param candidate: a string representing an extracted paragraph that
         we want to check to see if it matches a Knowledge Area.
    """

    for knowledgearea in knowledgeAreas:
        if knowledgearea.getText().strip() == str(candidate).strip():
            return knowledgearea

    return KnowledgeArea()


def ParagraphIsKnowledgeArea(pdocument: Document, p: Paragraph, i: int, lastPos: int, pText: str) -> Tuple[bool, str]:
    """
        Determine if the content of a paragraph consists of a Knowledge Area.
        Return a boolean and a string.

        :param pdocument: a Python Docx Document object containing the document we are parsing.
        :param p: a Python Docx Paragraph object containing the paragraph under examination.
        :param i: the integer representing a position in the enumeration of the paragraphs.
        :param lasPos: an integer representing the position of the last character in the paragraph.
        :param pText: a string containing the text from the paragraph.

        :return a Tuple containing the boolean value and the string representing the Knowledge Area.
    """

    global firstPositionOfCKA

    if p.style.name == constant.STYLE_KNOWLEDGEAREA and pText.strip() and pText[lastPos-1].isdigit():
        pPrev = document.paragraphs[i - 1]
        pPrevText = pPrev.text.strip()
        pPrevLasPos = len(pPrevText)

        pPrevRunCount = len(pPrev.runs)

        pPrevIsBold = 'false'
        if pPrevRunCount > 0:
            pPrevIsBold = 'true' if pPrev.runs[0].bold else 'false'

        if pPrevIsBold == 'false':
            if pPrevRunCount > 1:
                pPrevIsBold = 'true' if pPrev.runs[1].bold else 'false'

        if pPrevText != "":
            if pPrev.style.name == constant.STYLE_NORMAL and pPrevIsBold == 'true' and not (pPrevText[pPrevLasPos-1].isdigit()):
                if firstPositionOfCKA == 0:
                    firstPositionOfCKA = i

                return True, pPrevText

    return False, ""


def GetExistingCatalogRecords(firebase: firebase, documentName: str, documentSemester: str, documentYear: str) -> None:
    """
        Get existing catalog records from database and insert into global collection.

        :param firebase: a firebase connection
        :param documentName: a string representing the catalog document name
        :param documentSemester: a string representation of the catalog academic semester
        :param documentYear: a string representation of the catalog year
    """

    global existingCatalogRecords

    obj_key_list = []

    result = firebase.get('/catalog', None)

    if result is None:
        return

    for i in result.keys():
        obj_key_list.append(i)

    for i in obj_key_list:
        catalog = Catalog()
        catalog.setId(i)
        catalog.setDocumentName(result[i]['document_name'])
        catalog.setSemester(result[i]['semester'])
        catalog.setYear(result[i]['year'])
        existingCatalogRecords.append(catalog)


def ExtractCatalogRecord(firebase: firebase, documentName: str, documentSemester: str, documentYear: str) -> Catalog:
    """
        Look if catalog record exists, if not - write it.

        :param firebase: a firebase connection
        :param documentName: a string representing the catalog document name
        :param documentSemester: a string representation of the catalog academic semester
        :param documentYear: a string representation of the catalog year

        :return Catalog object
    """

    catalog = Catalog()
    catalog.setDocumentName(documentName)
    catalog.setSemester(documentSemester)
    catalog.setYear(documentYear)

    existingKey = FindExistingCatalogId(catalog)

    if existingKey == "":
        newCatalogRecord = {
            'document_name': documentName,
            'semester': documentSemester,
            'year': documentYear
        }
        result = firebase.post('catalog', newCatalogRecord)
        catalog.setId(result.get("name"))
    else:
        catalog.setId(existingKey)

    return catalog


def FindExistingCatalogId(newCatalogRecord: Catalog) -> str:
    """
        Look for existing Catalog Record and get the Id.

        :param newCatalogRecord: a Catalog object

        :return string representation of Id or empty string if not found
    """

    global existingCatalogRecords

    for existingCatalogRecord in existingCatalogRecords:
        if existingCatalogRecord.getDocumentName() == newCatalogRecord.getDocumentName():
            return existingCatalogRecord.getId()

    return ""


def GetExistingKnowledgeAreas(firebase: firebase, document: Document) -> None:
    """
        Get existing KnowledgeAreas from the database and
        insert them into the global collection.

        :param firebase: a firebase connection
        :param document: a Python-Docx Document object

    """

    global existingKnowledgeAreas

    obj_key_list = []

    result = firebase.get('/knowledgearea', None)

    if result is None:
        return

    for i in result.keys():
        obj_key_list.append(i)

    for i in obj_key_list:
        knowledgeArea = KnowledgeArea()
        knowledgeArea.setText(result[i]['content'])
        knowledgeArea.setId(i)
        existingKnowledgeAreas.append(knowledgeArea)


def ExtractKnowledgeAreas(firebase: firebase, document: Document, knowledgeAreas: List[KnowledgeArea]) -> None:
    """
        Find the KnowledgeAreas in the document and
        insert them into a database table.

        :param firebase: a firebase connection
        :param document: a Python-Docx Document object
        :param knowledgeAreas: a List of KnowledgeArea objects

    """

    global firstPositionOfCKA
    firstPositionOfCKA = 0

    for i, p in enumerate(document.paragraphs):
        pText = p.text.strip()
        lastPos = len(pText)
        kaText = ""

        isKnowledgeArea, kaText = ParagraphIsKnowledgeArea(document, p, i, lastPos, pText)

        if isKnowledgeArea and kaText:
            knowledgeArea = KnowledgeArea()
            knowledgeArea.setText(kaText)
            knowledgeAreas.append(knowledgeArea)

        diff = (i - firstPositionOfCKA)
        if firstPositionOfCKA > 0 and diff > 120:
            break

    for knowledgeArea in knowledgeAreas:
        newKnowledgeArea = {
            'content': knowledgeArea.getText()
        }

        existingKey = FindExistingKnowledgeAreaId(knowledgeArea)

        if existingKey == "":
            result = firebase.post('knowledgearea', newKnowledgeArea)
            knowledgeArea.setId(result.get("name"))
        else:
            knowledgeArea.setId(existingKey)


def FindExistingKnowledgeAreaId(newKnowledgeArea: KnowledgeArea) -> str:
    """
        Find existing KnowledgeArea in the global collection.
        If found, get the Id.

        :param newKnowledgeArea: a KnowledgeArea object

        :return string representation of the Id

    """

    global existingKnowledgeAreas

    for existingKnowledgeArea in existingKnowledgeAreas:
        if existingKnowledgeArea.getText() == newKnowledgeArea.getText():
            return existingKnowledgeArea.getId()

    return ""


def ExtractCourseAndDescription(firebase: firebase, document: Document, knowledgeAreas: List[KnowledgeArea], courses: List[Course], catalogId: str) -> None:
    """
        Build a list of courses with Foreign Key back to KnowledgeArea
        and insert them into the database.

        :param firebase: a firebase connection
        :param document: a Python-Docx Document object
        :param knowledgeAreas: a List of KnowledgeArea objects
        :param courses: a List of Course objects
        :param catalogId: a string representation of the current Catalog Id

    """

    global currentKnowledgeArea

    knowledgeAreaId = ""
    knowledgeAreaTitle = ""
    candidateId = ""
    partialTitle = []
    fulltitle = ""

    global courseNameStyles

    for p in document.paragraphs:
        if p.style.name == constant.STYLE_NORMAL:
            currentKnowledgeArea = getMatchingKnowledgeArea(p.text.strip())
            candidateId = currentKnowledgeArea.getId()

            if candidateId != "":
                knowledgeAreaId = str(candidateId)
                knowledgeAreaTitle = currentKnowledgeArea.getText()
        else:
            knowledgeAreaId = ""

        if knowledgeAreaId != "" and p.text.strip() != knowledgeAreaTitle and p.text.strip() != "":
            pText = p.text.strip()
            lastPos = len(pText)

            if not pText[lastPos-1].isdigit():
                partialTitle.append(p.text.strip())
            else:
                if len(partialTitle) > 0:
                    partialTitle.append(p.text.strip())
                    fulltitle = ' '.join(partialTitle)
                else:
                    fulltitle = pText

                fulltitle = fulltitle.replace("*", "")

                course = Course()
                course.setKnowledgeAreaId(knowledgeAreaId)
                course.setKnowledgeArea(currentKnowledgeArea.getText())
                course.setTOCEntry(p.text.strip())
                course.setTitle(fulltitle.rstrip(string.digits).strip())
                courses.append(course)
                partialTitle = []

    ExtractCourseDescriptions(document, courses)

    for course in courses:
        if course.getDescription().strip() == "":
            courseTitle = course.getTitle().strip()
            courseTitle = courseTitle.replace("New!", "").replace("*", "").replace("'","").strip()

            pNumber = 0
            takeDescription = 0
            descriptionElements = []

            for p in document.paragraphs:
                pNumber += 1
                    
                if p.style.name in courseNameStyles and p.text != "":
                    pText = p.text.replace("New! ", "").replace("*", "").replace("'","").strip()
                    if pText.__contains__(courseTitle) or courseTitle.__contains__(pText):
                        takeDescription = 1
                    else:
                        takeDescription = 0

                if p.style.name == constant.STYLE_BODY_TEXT and takeDescription == 1:
                    #print("taking {} as course description for course: {}".format(p.text.strip(), course.getTitle()))
                    descriptionElements.append(p.text.strip())

            fullDescription = course.getDescription()

            if len(descriptionElements) > 0:
                fullDescription = ' '.join(descriptionElements)

            course.setDescription(fullDescription)

    for course in courses:
        if course.getDescription().strip() == "":
            print("course: {} has no description".format(course.getTitle()))

        if course.getKnowledgeArea() == course.getTitle():
            continue

        newCourse = {
            'catalogid': catalogId,
            'knowledgeareaid': course.getKnowledgeAreaId(),
            'name': course.getTitle(),
            'description': course.getDescription()
        }
        result = firebase.post('course', newCourse)
        course.setId(result)


def HeaderIsCourseTitle(candidate: str) -> bool:
    """
        Compare candidate string with every course title until match is
        found or end of list.

        :param candidate: a string representation of a candidate course title

        :return boolean

    """

    candidate = candidate.replace("New!", "").replace("*", "").strip()
    candidate = candidate.replace("'","").strip()


    for course in courses:

        compareThis = []
        courseTitle = course.getTitle().strip()
        courseTitle = courseTitle.replace("New!", "").replace("*", "").strip()
        courseTitle = courseTitle.replace("'","").strip()

        lastPos = len(courseTitle)

        if courseTitle != "":

            if str(courseTitle) == str(candidate):
                return True

            if courseTitle.__contains__(candidate) or candidate.__contains__(courseTitle):
                return True

            if courseTitle.__contains__(":") and courseTitle[lastPos-1] != ":":
                (key, value) = courseTitle.split(':', 1)
                compareThis.append(key.strip())
                compareThis.append(value.strip())
                compareThis.append(courseTitle)
            else:
                compareThis.append(courseTitle)

            if any(ele in candidate for ele in compareThis) and len(candidate) >= len(courseTitle):
                return True
            else:
                if len(candidate) >= len(courseTitle):
                    for titlePart in compareThis:
                        if titlePart in candidate:
                            return True
                        if candidate in titlePart:
                            return True

            courseTitlePostCol = ""
            candidatePostCol = ""
            lastCandidatePos = len(candidate)

            if courseTitle.__contains__(":") and courseTitle[lastPos-1] != ":":
                (key, value) = courseTitle.split(':', 1)
                courseTitlePostCol =value.strip()

            if candidate.__contains__(":") and candidate[lastCandidatePos-1] != ":":
                (key, value) = candidate.split(':', 1)
                candidatePostCol =value.strip()

            if candidatePostCol == courseTitlePostCol:
                return True

    return False


def GetAssociatedCourse(courses: List[Course], candidate: str) -> Course:
    """
        Compare candidate string with every course title until match is
        found or end of list.

        :param candidate: a string representation of a candidate course title

        :return Course object

    """

    maybeCourseTitle = ""
    maybeCourseTitle = candidate.replace("*","").replace("'","").strip()

    for course in courses:

        compareThis = []
        courseTitle = course.getTitle().strip()
        courseTitle = courseTitle.replace("New!", "").replace("*", "").replace("'","").strip()
        lastPos = len(courseTitle)

        if courseTitle != "":

            #if courseTitle.lower().__contains__("culinary"):
                #print("inside GetAssociatedCourse with candidate: {} and courseTitle: {}".format(maybeCourseTitle.replace("'",""), courseTitle))
            
            if maybeCourseTitle.lower().replace("'", "") == courseTitle.lower().replace("'", ""):
                return course
            
            if courseTitle.__contains__(":") and courseTitle[lastPos-1] != ":":
                (key, value) = courseTitle.split(':', 1)
                compareThis.append(key.strip())
                compareThis.append(value.strip())
                compareThis.append(courseTitle)
            else:
                compareThis.append(courseTitle)

            if any(ele in maybeCourseTitle for ele in compareThis):
                return course

            courseTitlePostCol = ""
            candidatePostCol = ""
            lastCandidatePos = len(candidate)

            if courseTitle.__contains__(":") and courseTitle[lastPos-1] != ":":
                (key, value) = courseTitle.split(':', 1)
                courseTitlePostCol =value.strip()

            if candidate.__contains__(":") and candidate[lastCandidatePos-1] != ":":
                (key, value) = candidate.split(':', 1)
                candidatePostCol =value.strip()

            if candidatePostCol == courseTitlePostCol:
                return course                


    nullCourse = Course()
    nullCourse.setKnowledgeArea = "No Match"
    nullCourse.setTOCEntry = "No Match"
    nullCourse.setTitle = "No Match"
    return nullCourse


def ExtractCourseDescriptions(document: Document, courses: List[Course]) -> None:
    """
        Extract the course descriptions from the document.
        Write them to the appropriate Course records in the
        database.

        :param document: a Python-Docx Document object
        :param courses: a List of Course objects

    """

    courseDescriptionOn = False
    allparagraphs = []
    pertinentParagraphs = []
    number = 0
    pHeader = ""
    hIsCourseTitle = False
    partialTitle = []

    global courseNameStyles

    for p in document.paragraphs:

        if p.text.strip() != "":

            pText = p.text.strip()
            lastPos = len(pText)

            if pText[lastPos-1] == ":":
                partialTitle.append(pText)
            else:
                if len(partialTitle) > 0:
                    partialTitle.append(p.text.strip())

            number = number + 1
            paragraph = Paragraph()
            paragraph.setNumber(number)
            paragraph.setStyle(p.style.name)
            paragraph.setText(p.text.strip().rstrip(string.digits))
            allparagraphs.append(paragraph)

    for i, item in enumerate(allparagraphs):
        if item.getStyle() == constant.STYLE_BODY_TEXT:
            pPrev = allparagraphs[i-1]

            pHeader = ""
            if pPrev.getStyle() in courseNameStyles:
                number = number + 1
                p2 = Paragraph()
                p2.setNumber(number)
                p2.setStyle(pPrev.getStyle())
                p2.setText(pPrev.getText())
                paragraphs2.append(p2)
                pHeader = pPrev.getText().strip()
                if pHeader != "":
                    hIsCourseTitle = HeaderIsCourseTitle(pHeader)

            if hIsCourseTitle:
                number = number + 1
                pc = Paragraph()
                pc.setNumber(number)
                pc.setStyle(item.getStyle())
                pc.setHeader(pHeader)
                pc.setText(item.getText())
                pertinentParagraphs.append(pc)

    associatedCourse = Course()
    currentCourseDescription = ""
    courseTitlePosition = -1

    for i, p in enumerate(pertinentParagraphs):
        #if p.getHeader().strip() != "":
            #print("paragraph header: {}".format(p.getHeader().strip()))
            
        if p.getHeader().strip() != "" and courseDescriptionOn and i > courseTitlePosition:
            associatedCourse.setDescription(currentCourseDescription)
            courseDescriptionOn = False
            currentCourseDescription = ""
            courseTitlePosition = -1
        if p.getHeader().strip() != "" and not courseDescriptionOn:
            courseDescriptionOn = True
            associatedCourse = GetAssociatedCourse(courses, p.getHeader().replace("'","").strip())
            courseTitlePosition = i
            currentCourseDescription += p.getText().strip()
        if courseDescriptionOn and i > courseTitlePosition:
            currentCourseDescription += p.getText().strip()


def SentenceIsLO(sentence: str) -> Tuple[bool, str]:
    """
        Determine if the sentence meets the criteria for a learning
        outcome.

        :param sentence: a sentence from a course description

        :return a Tuple[boolean, string]

    """


    nullSentence = ""

    for word, pos in sentence.tags:
        if constant.LL_VERBS.count(word) >= 1:
            return 'true', str(sentence)
        if constant.LL_NOUNS.count(word) >= 1:
            return 'true', str(sentence)
        if constant.LL_ADJECTIVES.count(word) >= 1:
            return 'true', str(sentence)
        for phrase in constant.LL_PHRASES:
            if sentence.noun_phrases.count(phrase) >= 1:
                return 'true', str(sentence)

    return False, nullSentence


def ExtractLearningOutcomes(firebase: firebase, document: Document, courses: List[Course]) -> None:
    """
        Load each course description into BLOB and:
            - get each sentence,
            - tag parts of speech in each sentence,
            - compare verb with Bloom Action Verbs,
            - if match, select sentence as learning outcome,
            - insert into global collection.

        :param firebase: a firebase connection object
        :param document: a Python-Docx Document object
        :param courses: a List of Course objects

    """

    global learningObjectives
    courseId = ""

    for course in courses:

        courseId = course.getId()

        blob = TextBlob(course.getDescription())
        for sentence in blob.sentences:
            isLO, candidateLO = SentenceIsLO(sentence)

            if isLO:
                lo = LearningObjective()
                lo.setCourseId(courseId.get("name"))
                lo.setText(candidateLO)
                lo.setKnowledgeAreaId(course.getKnowledgeAreaId())
                learningObjectives.append(lo)


def WriteLearningOutcomes(firebase: firebase, learningObjectives: List[LearningObjective]) -> None:
    """

        Write the LearningObjectives captured in the global list
        to the database.

        :param firebase: firebase Database Connection
        :param learningObjectives: LearningObjective object list

    """

    for learningObjective in learningObjectives:
        newLearningObjective = {
            'courseid': learningObjective.getCourseId(),
            'content': learningObjective.getText(),
            'knowledgeareaid': learningObjective.getKnowledgeAreaId()
        }
        result = firebase.post('learningobjective', newLearningObjective)
        learningObjective.setId(result.get("name"))
        

def WriteLearningOutcome_Course_Records(firebase: firebase, learningObjectives: List[LearningObjective]) -> None:
    """
        
        Write the learningobjective_course bridge table
        
        :param firebase: firebase Database Connection
        :param learningObjectives: LearningObjective object list

    """

    for learningObjective in learningObjectives:
        newLearningObjective_Course = {
            'courseid': learningObjective.getCourseId(),
            'learningobjectiveid': learningObjective.getId()
        }
        firebase.post('learningobjective_course', newLearningObjective_Course)
    
    
    
if __name__ == "__main__":

    tStart = time.localtime()
    start_time = time.strftime("%H:%M:%S", tStart)
    print("Current Time: {}".format(start_time))
    t0 = time.time()


    print("Check database for existing Catalog Records.")
    #GetExistingCatalogRecords(firebase, documentName, documentSemester, documentYear)

    print("Extract Catalog Record from current document.")
    catalog = ExtractCatalogRecord(firebase, documentName, documentSemester, documentYear)
    catalogId = catalog.getId()

    print("Check database for existing Knowledge Areas.")
    #GetExistingKnowledgeAreas(firebase, document)

    print("Extract Knowledge Areas from the current document.")
    ExtractKnowledgeAreas(firebase, document, knowledgeAreas)

    print("Extract Courses and Descriptions from the current document.")
    ExtractCourseAndDescription(firebase, document, knowledgeAreas, courses, catalogId)

    print("Extract stated Learning Objectives from the current document.")
    ExtractLearningOutcomes(firebase, document, courses)

    print("Write the extracted Learning Objectives to the Database.")
    WriteLearningOutcomes(firebase, learningObjectives)

    print("Write the LearningObjective_Course bridge table.")
    WriteLearningOutcome_Course_Records(firebase, learningObjectives)
    
    tEnd = time.localtime()
    end_time = time.strftime("%H:%M:%S", tEnd)
    print("Current Time: {}".format(end_time))

    t1 = time.time()

    tElapsed = t1 - t0

    print("Time elapsed in seconds: {} ".format(tElapsed))


