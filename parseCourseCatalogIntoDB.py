import csv
import constant
from docx import Document
from textblob import TextBlob
from course import Course
from knowledgearea import KnowledgeArea
from catalog import Catalog
import string
from paragraph import Paragraph
import nltk
from nltk.stem.porter import PorterStemmer
from firebase import firebase
from docx.oxml.ns import qn
import xml.etree.ElementTree as ET
import json
import re
from typing import List, Set, Dict, Tuple, Optional

existingCatalogRecords = []
existingKnowledgeAreas = []        
firstPositionOfCKA = 0
currentKnowledgeArea = KnowledgeArea()
courseNameStyles = [constant.STYLE_NORMAL, constant.STYLE_HEADING]


firebase = firebase.FirebaseApplication('https://scholacity-org-test.firebaseio.com/')

document = Document('Fall2019_LLFullCatalog.docx')
documentName = "Fall2019_LLFullCatalog.docx";
documentSemester = "Fall";
documentYear = "2019";

# document = Document('Spring2020_LeisureLearningCatalogFULL.docx')
# documentName = "Spring2020_LeisureLearningCatalogFULL.docx";
# documentSemester = "Spring";
# documentYear = "2020";


knowledgeAreas = []
courses = []
paragraphs1 = []
paragraphs2 = []


def getMatchingKnowledgeArea(candidate: str) -> KnowledgeArea:
    # find matching knowledge area and return Id
    
    for knowledgearea in knowledgeAreas:
        if knowledgearea.getText().strip() == str(candidate).strip():
            return knowledgearea
        
    return KnowledgeArea()


def ParagraphIsKnowledgeArea(pdocument: Document, p: Paragraph, i: int, lastPos: int, pText: str) -> Tuple[bool, str]:
    # make the determination that the content of a paragraph consists of a Knowledge Area
    
    global firstPositionOfCKA
    
    if p.style.name == constant.STYLE_KNOWLEDGEAREA and pText.strip() and pText[lastPos-1].isdigit():
        pPrev = document.paragraphs[i - 1]
        pPrevText = pPrev.text.strip()
        pPrevLasPos = len(pPrevText)
        
        pPrevRunCount = len(pPrev.runs)
        
        pPrevIsBold = 'false'
        if pPrevRunCount > 0:
            pPrevIsBold = 'true' if pPrev.runs[0].bold else 'false' 
            
        if pPrevIsBold == 'false':
            if pPrevRunCount > 1:
                pPrevIsBold = 'true' if pPrev.runs[1].bold else 'false' 
        
        
        if pPrevText != "":
            if pPrev.style.name == constant.STYLE_NORMAL and pPrevIsBold == 'true' and not (pPrevText[pPrevLasPos-1].isdigit()):
                if firstPositionOfCKA == 0:
                    firstPositionOfCKA = i
        
                return True, pPrevText
            
    return False, ""


def GetExistingCatalogRecords(firebase: firebase, documentName: str, documentSemester: str, documentYear: str) -> None:
    # get existing catalog records from database and put into array
    
    global existingCatalogRecords
    
    obj_key_list = []
    
    result = firebase.get('/Catalog', None)
    
    if result is None:
        return
    
    for i in result.keys():
        obj_key_list.append(i)
        
    
    for i in obj_key_list:
        catalog = Catalog()
        catalog.setId(i)
        catalog.setDocumentName(result[i]['document_name'])
        catalog.setSemester(result[i]['semester'])
        catalog.setYear(result[i]['year'])
        existingCatalogRecords.append(catalog)
   
    
    
def ExtractCatalogRecord(firebase: firebase, documentName: str, documentSemester: str, documentYear: str) -> Catalog:
    # look if catalog record exists, if not - write it

   catalog = Catalog()
   catalog.setDocumentName(documentName)
   catalog.setSemester(documentSemester)
   catalog.setYear(documentYear)
   
   existingKey = FindExistingCatalogId(catalog)

   if existingKey == "":
       newCatalogRecord = {
            'document_name': documentName,
            'semester': documentSemester,
            'year': documentYear
       }
       result = firebase.post('Catalog', newCatalogRecord)
       catalog.setId(result.get("name"))
   else:
       catalog.setId(existingKey)
   
    
   return catalog


   
def FindExistingCatalogId(newCatalogRecord: Catalog) -> str:
    # find existing knowledgeArea - if found return Id
    
    global existingCatalogRecords
    
    for existingCatalogRecord in existingCatalogRecords:
        if existingCatalogRecord.getDocumentName() == newCatalogRecord.getDocumentName():
            return existingCatalogRecord.getId()
        
    return ""


def GetExistingKnowledgeAreas(firebase: firebase, document: Document) -> None:
    # get existing knowledge areas from the database and put into array
    
    global existingKnowledgeAreas
    
    obj_key_list = []
    
    result = firebase.get('/KnowledgeArea', None)
    
    if result is None:
        return
    
    for i in result.keys():
        obj_key_list.append(i)
        
    
    
    for i in obj_key_list:
        knowledgeArea = KnowledgeArea()
        knowledgeArea.setText(result[i]['Content'])
        knowledgeArea.setId(i)
        existingKnowledgeAreas.append(knowledgeArea) 
   
    
def ExtractKnowledgeAreas(firebase: firebase, document: Document, knowledgeAreas: List[KnowledgeArea]) -> None:
    # find the knowledge areas in the document and write them to a database table
    
    global firstPositionOfCKA
    firstPositionOfCKA = 0
    
    for i, p in enumerate(document.paragraphs):
        pText = p.text.strip()
        lastPos = len(pText)
        kaText = ""
        
        isKnowledgeArea, kaText = ParagraphIsKnowledgeArea(document, p, i, lastPos, pText)
        
        
        if isKnowledgeArea and kaText:
            knowledgeArea = KnowledgeArea()
            knowledgeArea.setText(kaText)
            knowledgeAreas.append(knowledgeArea)
                    
                    
        diff = (i - firstPositionOfCKA)
        if firstPositionOfCKA > 0 and diff > 120:
            break;            
                
    for knowledgeArea in knowledgeAreas:
        newKnowledgeArea = {
            'Content':knowledgeArea.getText()    
        }

        existingKey = FindExistingKnowledgeAreaId(knowledgeArea)
        
        if existingKey == "":
            result = firebase.post('KnowledgeArea',newKnowledgeArea)
            knowledgeArea.setId(result.get("name"))
        else:
            knowledgeArea.setId(existingKey)


def FindExistingKnowledgeAreaId(newKnowledgeArea: KnowledgeArea) -> str:
    # find existing knowledgeArea - if found return Id
    
    global existingKnowledgeAreas
    
    for existingKnowledgeArea in existingKnowledgeAreas:
        if existingKnowledgeArea.getText() == newKnowledgeArea.getText():
            return existingKnowledgeArea.getId()
        
    return ""
    
        
def ExtractCourseAndDescription(firebase: firebase, document: Document, knowledgeAreas: List[KnowledgeArea], courses: List[Course], catalogId: str) -> None:
    # build a list of courses with foreign key back to knowledge area and insert into DB
    
    global currentKnowledgeArea
    knowledgeAreaId = ""
    knowledgeAreaTitle = ""
    candidateId = ""
    partialTitle = []
    fulltitle = ""
    
    global courseNameStyles
   
    for p in document.paragraphs:
        if p.style.name == constant.STYLE_NORMAL:
            currentKnowledgeArea = getMatchingKnowledgeArea(p.text.strip())
            candidateId = currentKnowledgeArea.getId()
            
            if candidateId != "":
                knowledgeAreaId = str(candidateId)
                knowledgeAreaTitle = currentKnowledgeArea.getText()
                print("knowledgeAreaId:{}".format(knowledgeAreaId))
                  
        else:
            knowledgeAreaId = ""
    
       
        if knowledgeAreaId != "" and p.text.strip() != knowledgeAreaTitle and p.text.strip() != "":
            pText = p.text.strip()
            lastPos = len(pText)

            if not pText[lastPos-1].isdigit():
                partialTitle.append(p.text.strip())
            else:
                if len(partialTitle) > 0:
                    partialTitle.append(p.text.strip())
                    fulltitle = ' '.join(partialTitle)
                else:
                    fulltitle = pText
                    
                  
                fulltitle = fulltitle.replace("*","")
                    
                course = Course()
                course.setKnowledgeAreaId(knowledgeAreaId)
                course.setKnowledgeArea(currentKnowledgeArea.getText())
                course.setTOCEntry(p.text.strip())
                course.setTitle(fulltitle.rstrip(string.digits).strip())
                courses.append(course)
                partialTitle = []

    ExtractCourseDescriptions(document, courses)   
    
    
    for course in courses:
        if course.getDescription() == "":
            pNumber = 0
            takeDescription = 0
            descriptionElements = []
            
            for p in document.paragraphs:
                pNumber += 1
                #if p.style.name == constant.STYLE_NORMAL and p.text != "" and p.runs[0].font.name == "Calibri" and p.runs[0].font.bold:
                if p.style.name in courseNameStyles and p.text != "":
                    pText = p.text.replace("New! ","").replace("*","").strip()
                    if p.text.__contains__(course.getTitle()) or course.getTitle().__contains__(pText):
                        #print("takeDescription is on")
                        takeDescription = 1
                    else:
                        if takeDescription == 1 and p.text.strip().__contains__("Sessions"):
                            takeDescription = 1
                        if takeDescription == 1 and p.text.strip().__contains__("Magic"):
                            takeDescription = 1
                        else:
                            takeDescription = 0
                        
                if p.style.name == constant.STYLE_BODY_TEXT and takeDescription == 1:
                    descriptionElements.append(p.text.strip())
                    
            fullDescription = course.getDescription()
            
            if len(descriptionElements) > 0:
                fullDescription = ' '.join(descriptionElements)
                
            course.setDescription(fullDescription)
            
    #fullText = []
    # counter = 0
    # for t in document.tables:
    #     counter += 1
    #     for row in t.rows:
    #         for cell in row.cells:
    #             for paragraph in cell.paragraphs:
    #                 if counter < 21:
    #                     #print("paragraph in table: {}".format(paragraph.text))
    #     if counter >= 20:
    #         break;
            
        #print("Table: {}".format('\n'.join(fullText)))
            
                
    
    for course in courses:
        if course.getKnowledgeArea() == course.getTitle():
            continue
        
        newCourse = {
            'CatalogId': catalogId,
            'KnowledgeAreaId':course.getKnowledgeAreaId(),
            'Name':course.getTitle(),
            'Description':course.getDescription()
        }
        result = firebase.post('Course',newCourse)
        course.setId(result)


def ExtractCourses(firebase: firebase, document: Document, knowledgeAreas: List[KnowledgeArea], courses: List[Course]) -> None:
    # build a list of courses with foreign key back to knowledge area and insert into DB
    
    global currentKnowledgeArea
    knowledgeAreaId = ""
    knowledgeAreaTitle = ""
    candidateId = ""
    partialTitle = []
    fulltitle = ""
   
    for p in document.paragraphs:
        if p.style.name == constant.STYLE_NORMAL:
            currentKnowledgeArea = getMatchingKnowledgeArea(p.text.strip())
            candidateId = currentKnowledgeArea.getId()
            
            if candidateId != "":
                knowledgeAreaId = str(candidateId)
                knowledgeAreaTitle = currentKnowledgeArea.getText()
                #print("knowledgeAreaId:{}".format(knowledgeAreaId))
                  
        else:
            knowledgeAreaId = ""
    
       
        if knowledgeAreaId != "" and p.text.strip() != knowledgeAreaTitle and p.text.strip() != "":
            pText = p.text.strip()
            lastPos = len(pText)

            if not pText[lastPos-1].isdigit():
                partialTitle.append(p.text.strip())
            else:
                if len(partialTitle) > 0:
                    partialTitle.append(p.text.strip())
                    fulltitle = ' '.join(partialTitle)
                else:
                    fulltitle = pText
                    
                  
                fulltitle = fulltitle.replace("*","")
                    
                course = Course()
                course.setKnowledgeAreaId(knowledgeAreaId)
                course.setKnowledgeArea(currentKnowledgeArea.getText())
                course.setTOCEntry(p.text.strip())
                course.setTitle(fulltitle.rstrip(string.digits).strip())
                courses.append(course)
                partialTitle = []

    ExtractCourseDescriptions(document, courses)   
    
    for course in courses:
        if course.getKnowledgeArea() == course.getTitle():
            continue
        
        newCourse = {
            'KnowledgeAreaId':course.getKnowledgeAreaId(),
            'Name':course.getTitle(),
            'Description':course.getDescription()
        }
        result = firebase.post('Course',newCourse)
        course.setId(result)
        


def HeaderIsCourseTitle(candidate: str) -> bool:
    # compare candidate with every course title until match found or eol    

    candidate = candidate.replace("New!","").replace("*","").strip()
    
    for course in courses:
        
        compareThis = []
        courseTitle = course.getTitle().strip()
        lastPos = len(courseTitle)
        #lastChar = courseTitle.charAt(lastPos-1)
        
        #print("course title: {}".format(courseTitle))
           
        if courseTitle != "":
            
            if str(courseTitle) == str(candidate):
                return True
            
            if courseTitle.__contains__(candidate) or candidate.__contains__(courseTitle):
                return True
            
            if courseTitle.__contains__(":") and courseTitle[lastPos-1] != ":":
                (key,value) = courseTitle.split(':',1)
                compareThis.append(key.strip())
                compareThis.append(value.strip())
                compareThis.append(courseTitle)
            else:
                compareThis.append(courseTitle)
            
       
            if any(ele in candidate for ele in compareThis) and len(candidate) >= len(courseTitle):
                return True
            else:
                if len(candidate) >= len(courseTitle):
                    for titlePart in compareThis:
                        if titlePart in candidate:
                            return True
                    
    return False    
        
def GetAssociatedCourse(courses: List[Course],candidate: str) -> Course:
    # compare candidate with every course title until match found or eol    
    
    
    candidate = candidate.replace("New! ","").replace("*","").strip()
        
        
    for course in courses:
        #print("course title: {}".format(course.getTitle()))
        
        compareThis = []
        courseTitle = course.getTitle().strip()
        lastPos = len(courseTitle)
        #lastChar = courseTitle.charAt(lastPos-1)
        
        
        if courseTitle != "":
            
            if courseTitle.__contains__(":") and courseTitle[lastPos-1] != ":":
                (key,value) = courseTitle.split(':',1)
                compareThis.append(key.strip())
                compareThis.append(value.strip())
                compareThis.append(courseTitle)
            else:
                compareThis.append(courseTitle)
            
       
            #if any(ele in candidate for ele in compareThis) and len(candidate) >= len(courseTitle):
            if any(ele in candidate for ele in compareThis):
                return course
        
    nullCourse = Course()
    nullCourse.setKnowledgeArea = "No Match"
    nullCourse.setTOCEntry = "No Match"
    nullCourse.setTitle = "No Match"
    return nullCourse    
    

def ExtractCourseDescriptions(document: Document, courses: List[Course]) -> None:
    
    courseDescriptionOn = False
    allparagraphs = []
    pertinentParagraphs = []
    number = 0
    pHeader = ""
    headerIsCourseTitle = False
    partialTitle = []
    fulltitle = ""
    
    global courseNameStyles
    
    for p in document.paragraphs:
        
        
        if p.text.strip() != "":
            
            pText = p.text.strip()
            lastPos = len(pText)

            if pText[lastPos-1] == ":":
                partialTitle.append(pText)
            else:
                if len(partialTitle) > 0:
                    partialTitle.append(p.text.strip())
                    fulltitle = ' '.join(partialTitle)
                else:
                    fulltitle = pText
                
            number = number + 1
            paragraph = Paragraph()
            paragraph.setNumber(number)
            paragraph.setStyle(p.style.name)
            paragraph.setText(p.text.strip().rstrip(string.digits))
            allparagraphs.append(paragraph)
            
            
    for i, item in enumerate(allparagraphs):
        if item.getStyle() == constant.STYLE_BODY_TEXT:
            pPrev = allparagraphs[i-1]
            
            pHeader = ""
            if pPrev.getStyle() in courseNameStyles:
                number = number + 1
                p2 = Paragraph()
                p2.setNumber(number)
                p2.setStyle(pPrev.getStyle())
                p2.setText(pPrev.getText())
                paragraphs2.append(p2)
                pHeader = pPrev.getText().strip()
                if pHeader != "":
                    headerIsCourseTitle = HeaderIsCourseTitle(pHeader)
                    
            if headerIsCourseTitle:
                number = number + 1
                pc = Paragraph()
                pc.setNumber(number)
                pc.setStyle(item.getStyle())
                pc.setHeader(pHeader)
                pc.setText(item.getText())
                pertinentParagraphs.append(pc)
            
    
    associatedCourse = Course()  
    currentCourseDescription = "" 
    courseTitlePosition = -1
     
    
    for i, p in enumerate(pertinentParagraphs):
        courseTitleCandidate = p.getHeader().strip()
        #print("courseTitleCandidate: {}".format(courseTitleCandidate))
        if p.getHeader().strip() != "" and courseDescriptionOn and i > courseTitlePosition:
            associatedCourse.setDescription(currentCourseDescription)
            courseDescriptionOn = False
            currentCourseDescription = ""
            courseTitlePosition = -1
        if p.getHeader().strip() != "" and not courseDescriptionOn:
            courseDescriptionOn = True
            associatedCourse = GetAssociatedCourse(courses,p.getHeader().strip())
            #print("return from getAssociatedCourse(): {}".format(associatedCourse.getTitle()))
            courseTitlePosition = i
            currentCourseDescription += p.getText().strip()
        if courseDescriptionOn and i > courseTitlePosition:
            currentCourseDescription += p.getText().strip()

        #print("currentCourseDescription: {}".format(currentCourseDescription))



def SentenceIsLO(sentence: str) -> Tuple[bool, str]:
    # determine if the sentence meets the criteria for a learning outcome
    
    nullSentence = ""
    
    for word, pos in sentence.tags:
        if constant.LL_VERBS.count(word) >= 1:
            return 'true', str(sentence)
        if constant.LL_NOUNS.count(word) >= 1:
            return 'true', str(sentence)
        if constant.LL_ADJECTIVES.count(word) >= 1:
            return 'true', str(sentence)
        for phrase in constant.LL_PHRASES:
            if sentence.noun_phrases.count(phrase) >= 1:
                return 'true', str(sentence)
            
    return False, nullSentence


def ExtractLearningOutcomes(firebase: firebase, document: Document, courses: List[Course]) -> None:
    # load each course description into blob and get the sentences
    # tag parts of speech in each sentence
    # compare verb in each sentence to list of Bloom action verbs
    # if match, select sentence as a learning outcome
    
    
    for course in courses:
        
        #print("course: {}".format(course.getTitle()))
       
        blob = TextBlob(course.getDescription())
        for sentence in blob.sentences:
            isLO, candidateLO = SentenceIsLO(sentence)
            
            if isLO:
                newLO = {
                    'CourseId':course.getId(),
                    'Text':candidateLO
                }
                result = firebase.post('LearningObjective', newLO)
                
   
       
GetExistingCatalogRecords(firebase, documentName, documentSemester, documentYear)

catalog = ExtractCatalogRecord(firebase, documentName, documentSemester, documentYear)
catalogId = catalog.getId()

GetExistingKnowledgeAreas(firebase, document)
ExtractKnowledgeAreas(firebase, document, knowledgeAreas)
ExtractCourseAndDescription(firebase, document, knowledgeAreas, courses, catalogId)
ExtractLearningOutcomes(firebase, document, courses)

          
