#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Sun Jan 19 12:35:02 2020

@author: Steve Satterfield
"""

import csv
import constant
from docx import Document
from textblob import TextBlob
from course import Course
import string
from paragraph import Paragraph
import re
import xlsxwriter
import nltk
        
knowledgeAreas = []

allCourseFile = "Courses.csv"
allCourseFile2 = "CoursesWithDescriptionsxxx.csv"
allCourseFile3 = "CoursesWithDescriptionsPartial.csv"

courses = []
coursesWithDescriptions = []

paragraphs1 = []
paragraphs2 = []

number = 0
document = Document('Fall2019_LLFullCatalog.docx')

currentKnowledgeArea = ""

def writeDocStyles():
    global number
    writeFile = "docStyles3.csv"
    wrawFile = "docStylesRaw.csv"
    
    global document
    
    with open(writeFile, 'w') as output:
        fieldnames = ['number','style','text']
        writer = csv.DictWriter(output, fieldnames=fieldnames)
    
        writer.writeheader()
        for p in document.paragraphs:
            if p.text.strip() != "":
                number = number + 1
                paragraph = Paragraph()
                paragraph.setNumber(number)
                paragraph.setStyle(p.style.name)
                paragraph.setText(p.text.strip().replace("New!","").rstrip(string.digits).lstrip())
                paragraphs1.append(paragraph)
                
        number = 0 
        i = 0
        
        for i, item in enumerate(paragraphs1):
            if item.getStyle() == constant.STYLE_BODY_TEXT:
                pPrev = paragraphs1[i-1]
                if pPrev.getStyle() == constant.STYLE_NORMAL:
                    number = number + 1
                    p2 = Paragraph()
                    p2.setNumber(number)
                    p2.setStyle(pPrev.getStyle())
                    p2.setText(pPrev.getText())
                    paragraphs2.append(p2)
                number = number + 1
                pc = Paragraph()
                pc.setNumber(number)
                pc.setStyle(item.getStyle())
                pc.setText(item.getText().strip())
                paragraphs2.append(pc)
                
                
        for paragraph in paragraphs2:
            writer.writerow({'number': paragraph.getNumber(), 'style': paragraph.getStyle(), 'text': paragraph.getText()})  
            
    with open(wrawFile, "w") as rawoutput:
        fieldnames = ['number','style','text']
    
        rwriter = csv.DictWriter(rawoutput, fieldnames=fieldnames)
    
        rwriter.writeheader()

        for paragraph in paragraphs1:
            rwriter.writerow({'number': paragraph.getNumber(), 'style': paragraph.getStyle(), 'text': paragraph.getText()})  
 
def writeTables():
    # find and output the tables in the document
           
    global document
    writeFile = "docTables.csv"

    with open(writeFile, 'w') as outputTables:
        fieldnames = ['table values']
        tableWriter = csv.writer(outputTables)

        for i, p in enumerate(document.tables):
            numRows = len(p.rows)
            numColumns = len(p.columns)
            
            print("table {}: has {} rows".format(i, numRows))
            print("table {}: has {} columns".format(i, numColumns))
            for x in range(numRows):
                for y in range(numColumns):
                    thisCell = p.cell(x,y)
                    print("cell[{},{}] contents is {}".format(x,y,thisCell.text))
                    tableWriter.writerow({"table {}: cell[{},{}] content is {}".format(i,x,y,thisCell.text)})
            
                
    
def writeKnowledgeAreas(ws):
    # find and output the knowledge areas
    
    global document
    firstPositionOfCKA = 0
    global knowledgeAreas
    
    row = 0
    col = 0
    
    ws.write(row,col, 'Knowledge Area')
    row += 1
    
    for i, p in enumerate(document.paragraphs):
        pText = p.text.strip()
        lastPos = len(pText)
        
        if p.style.name == constant.STYLE_NORMAL and pText != "" and pText[lastPos-1].isdigit():
            pPrev = document.paragraphs[i - 1]
            pPrevText = pPrev.text.strip()
            pPrevLasPos = len(pPrevText)
            
            pPrevRunCount = len(pPrev.runs)
            pPrevIsBold = 'false'
            if pPrevRunCount > 0:
                pPrevIsBold = 'true' if pPrev.runs[0].bold else 'false' 
                
            if pPrevIsBold == 'false':
                if pPrevRunCount > 1:
                    pPrevIsBold = 'true' if pPrev.runs[1].bold else 'false' 
            
            if pPrevText != "":
                if pPrev.style.name == constant.STYLE_NORMAL and pPrevIsBold == 'true' and not (pPrevText[pPrevLasPos-1].isdigit()):
                    if firstPositionOfCKA == 0:
                        firstPositionOfCKA = i
                    #print("candidate KnowledgeArea: {} is bold: {}".format(pPrevText, pPrevIsBold))
                    knowledgeAreas.append(pPrevText)
                    ws.write(row, col, pPrevText)
                    row += 1
                    
                    
        #print("position: {} - firstPositionOfCKA: {} -> {}".format(i, firstPositionOfCKA, (i - firstPositionOfCKA)))
        diff = (i - firstPositionOfCKA)
        if firstPositionOfCKA > 0 and diff > 120:
            break;            
                
    ws.set_column(0, 0, 35, default_format)
       
            
            
def writeCourses(ws):
    # build a list of courses and output a worksheet 
    global document
    global currentKnowledgeArea
    global courses
    global default_format

    for p in document.paragraphs:
         
       #print(p.text.strip())
        if p.style.name == constant.STYLE_NORMAL:
            if knowledgeAreas.count(p.text.strip()) == 1:
                currentKnowledgeArea = knowledgeAreas[knowledgeAreas.index(p.text.strip())]
            else:
                currentKnowledgeArea = currentKnowledgeArea
        else:
            currentKnowledgeArea = ""
    
       
        if currentKnowledgeArea != "" and p.text.strip() != "":
            course = Course()
            course.setKnowledgeArea(currentKnowledgeArea)
            course.setTOCEntry(p.text.strip())
            course.setTitle(p.text.strip().rstrip(string.digits))
            courses.append(course)

    row = 0
    col = 0
    ws.write(row, col, 'Knowledge Area')
    col += 1
    ws.write(row, col, 'Course Title')
    col = 0
    row += 1
    
    for course in courses:
        if course.getKnowledgeArea() != course.getTitle():
            ws.write(row, col, course.getKnowledgeArea())
            col += 1
            ws.write(row, col, course.getTitle())
            col += 1
            ws.write(row, col, course.getDescription())
            col = 0
            row += 1

    ws.set_column(0, 0, 35, default_format)
    ws.set_column(1, 1, 50, default_format)


def HeaderIsCourseTitle(candidate):
    # compare candidate with every course title until match found or eol    
    
    for course in courses:
        #print("course title: {}".format(course.getTitle()))
        
        compareThis = []
        courseTitle = course.getTitle().strip()
        lastPos = len(courseTitle)
        #lastChar = courseTitle.charAt(lastPos-1)
        
        
        if courseTitle != "":
            
            if courseTitle.__contains__(":") and courseTitle[lastPos-1] != ":":
                (key,value) = courseTitle.split(':',1)
                compareThis.append(key.strip())
                compareThis.append(value.strip())
                compareThis.append(courseTitle)
            else:
                compareThis.append(courseTitle)
            
       
            if any(ele in candidate for ele in compareThis) and len(candidate) >= len(courseTitle):
                return 'true'
        
    return 'false'    
        
def GetAssociatedCourse(candidate):
    # compare candidate with every course title until match found or eol    
    
    for course in courses:
        #print("course title: {}".format(course.getTitle()))
        
        compareThis = []
        courseTitle = course.getTitle().strip()
        lastPos = len(courseTitle)
        #lastChar = courseTitle.charAt(lastPos-1)
        
        
        if courseTitle != "":
            
            if courseTitle.__contains__(":") and courseTitle[lastPos-1] != ":":
                (key,value) = courseTitle.split(':',1)
                compareThis.append(key.strip())
                compareThis.append(value.strip())
                compareThis.append(courseTitle)
            else:
                compareThis.append(courseTitle)
            
       
            if any(ele in candidate for ele in compareThis) and len(candidate) >= len(courseTitle):
                return course
        
    nullCourse = Course()
    nullCourse.setKnowledgeArea = "No Match"
    nullCourse.setTOCEntry = "No Match"
    nullCourse.setTitle = "No Match"
    return nullCourse    
    
def writeLearningOutcomes(ws):
    # load each course description into blob and get the sentences
    # tag parts of speech in each sentence
    # compare verb in each sentence to list of Bloom action verbs
    # if match, select sentence as a learning outcome
    global coursesWithDescriptions
    global default_format
    global description_format
    
    posFile = "taggedwords.csv"
    
    row = 0
    col = 0
    ws.write(row, col, 'Knowledge Area')
    col += 1
    ws.write(row, col, 'Course Title')
    col += 1
    ws.write(row, col, 'Description')
    
    with open(posFile, 'w') as outputTables:
        posWriter = csv.writer(outputTables)

        for course in coursesWithDescriptions:
            col = 0
            row += 1
            ws.write(row, col, course.getKnowledgeArea())
            col += 1
            ws.write(row, col, course.getTitle())
            col += 1
            ws.write(row, col, course.getDescription())
            col += 1
            
            firstRowInCourse = row
            
            blob = TextBlob(course.getDescription())
            for sentence in blob.sentences:
                #print(sentence)
                for word, pos in sentence.tags:
                    posWriter.writerow({"word: {}, pos: {}, startswith V: {}".format(word, pos, pos.startswith("V"))})  
                    if constant.BLOOM_ACTION_WORDS.count(word) >= 1:
                        print(sentence)
                        print({"word: {}, pos: {}, startswith V: {}".format(word, pos, pos.startswith("V"))})
                        ws.write(row, col, str(sentence))
                        row += 1
                        break
                    
    ws.set_column(0, 0, 35, default_format)
    ws.set_column(1, 1, 50, default_format)
    ws.set_column(2, 2, 120, description_format)
    ws.set_column(3, 3, 120, description_format)

                    
    
                
            
    
def writeCoursesWithDescriptions(ws):
    
    global document
    global courses
    global coursesWithDescriptions
    courseDescriptionOn = 'false'
    allparagraphs = []
    pertinentParagraphs = []
    global number
    global default_format
    global description_format
    
    
    writeFile = "pertinentParagraphs.csv"
    pHeader = ""
    headerIsCourseTitle = 'false'
    
    for p in document.paragraphs:
        if p.text.strip() != "":
            number = number + 1
            paragraph = Paragraph()
            paragraph.setNumber(number)
            paragraph.setStyle(p.style.name)
            paragraph.setText(p.text.strip().rstrip(string.digits))
            allparagraphs.append(paragraph)
            
    for i, item in enumerate(allparagraphs):
        if item.getStyle() == constant.STYLE_BODY_TEXT:
            pPrev = allparagraphs[i-1]
            pHeader = ""
            if pPrev.getStyle() == constant.STYLE_NORMAL:
                number = number + 1
                p2 = Paragraph()
                p2.setNumber(number)
                p2.setStyle(pPrev.getStyle())
                p2.setText(pPrev.getText())
                paragraphs2.append(p2)
                pHeader = pPrev.getText().strip()
                if pHeader != "":
                    headerIsCourseTitle = HeaderIsCourseTitle(pHeader)
                    
            if headerIsCourseTitle == 'true':
                number = number + 1
                pc = Paragraph()
                pc.setNumber(number)
                pc.setStyle(item.getStyle())
                pc.setHeader(pHeader)
                pc.setText(item.getText())
                pertinentParagraphs.append(pc)
            
    
    associatedCourse = Course()  
    currentCourseDescription = "" 
    courseTitlePosition = -1
     
    row = 0
    col = 0
    ws.write(row, col, 'Knowledge Area')
    col += 1
    ws.write(row, col, 'Course Title')
    col += 1
    ws.write(row, col, 'Description')
    col = 0
    row += 1
    
    for i, p in enumerate(pertinentParagraphs):
        courseTitleCandidate = p.getHeader().strip()
        if p.getHeader().strip() != "" and courseDescriptionOn == 'true' and i > courseTitlePosition:
            newCourse = Course()
            newCourse.setKnowledgeArea(associatedCourse.getKnowledgeArea())
            newCourse.setTitle(associatedCourse.getTitle())
            newCourse.setDescription(currentCourseDescription)
            coursesWithDescriptions.append(newCourse)
            courseDescriptionOn = 'false'
            ws.write(row, col, associatedCourse.getKnowledgeArea())
            col += 1
            ws.write(row, col, associatedCourse.getTitle())
            col += 1
            ws.write(row, col, currentCourseDescription)
            col = 0
            row += 1
            
            currentCourseDescription = ""
            courseTitlePosition = -1
        if p.getHeader().strip() != "" and courseDescriptionOn == 'false':
            courseDescriptionOn = 'true'
            associatedCourse = GetAssociatedCourse(p.getHeader().strip())
            courseTitlePosition = i
            currentCourseDescription += p.getText().strip()
        if courseDescriptionOn == 'true' and i > courseTitlePosition:
            currentCourseDescription += p.getText().strip()

    ws.set_column(0, 0, 35, default_format)
    ws.set_column(1, 1, 50, default_format)
    ws.set_column(2, 2, 120, description_format)
    
       
                    

workbook = xlsxwriter.Workbook('Course_Catalog_Extraction.xlsx')
wsKnowledgeAreas = workbook.add_worksheet()
wsKnowledgeAreas.name = "Knowledge Areas"

wsCourses = workbook.add_worksheet()
wsCourses.name = "Courses"

wsCourseDescriptions = workbook.add_worksheet()
wsCourseDescriptions.name = "Course Descriptions"

wsLearningOutcomes = workbook.add_worksheet()
wsLearningOutcomes.name = "Stated Learning Outcomes"

default_format = workbook.add_format({'valign':'top'})
description_format = workbook.add_format({'text_wrap':'true', 'valign':'top'})

writeKnowledgeAreas(wsKnowledgeAreas)
writeCourses(wsCourses)
writeCoursesWithDescriptions(wsCourseDescriptions)
writeLearningOutcomes(wsLearningOutcomes)

workbook.close()

        
    
